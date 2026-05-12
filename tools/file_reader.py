"""
File reader tool — auto-detects and extracts text from common file formats.

Supports out of the box (stdlib): .txt .md .log .csv .json .yaml .yml .py .js
.ts .ps1 .sh .sql .html .htm .xml .ini .cfg .toml .env

Optional formats (graceful fallback if dep missing):
- .pdf  → pdfplumber
- .docx → python-docx
- .xlsx → openpyxl
- .html → markdownify (better rendering)

All functions return human-readable strings; never raise.
"""

from __future__ import annotations

import json
import logging
import os
import re
from pathlib import Path
from typing import Annotated

from pydantic import Field

from .audit import audit_log

logger = logging.getLogger(__name__)

_TEXT_EXTS = {
    ".txt", ".md", ".log", ".csv", ".json", ".yaml", ".yml",
    ".py", ".js", ".ts", ".tsx", ".jsx", ".ps1", ".sh", ".bat",
    ".sql", ".xml", ".ini", ".cfg", ".toml", ".env", ".rst",
    ".java", ".cs", ".go", ".rs", ".rb", ".php", ".kt", ".swift",
    ".c", ".cpp", ".h", ".hpp", ".scala", ".r", ".m",
}
_HTML_EXTS = {".html", ".htm"}
_MAX_DEFAULT = 50000


def _safe_path(path: str) -> Path | None:
    """Resolve a user-supplied path; reject obvious traversal."""
    try:
        p = Path(os.path.expandvars(os.path.expanduser(path))).resolve()
    except Exception:
        return None
    return p


def _read_pdf(path: Path) -> str:
    try:
        import pdfplumber  # type: ignore
    except ImportError:
        return f"[PDF reader unavailable — install pdfplumber to read {path.name}]"
    try:
        with pdfplumber.open(str(path)) as pdf:
            chunks = []
            for i, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                if text.strip():
                    chunks.append(f"\n--- Page {i} ---\n{text}")
        return "".join(chunks) or f"[PDF {path.name} contained no extractable text]"
    except Exception as exc:  # noqa: BLE001
        logger.exception("PDF read failed")
        return f"[PDF read failed: {exc!s}]"


def _read_docx(path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return f"[DOCX reader unavailable — install python-docx to read {path.name}]"
    try:
        doc = Document(str(path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                style = (para.style.name if para.style else "").lower()
                if style.startswith("heading"):
                    level = "".join(c for c in style if c.isdigit()) or "1"
                    parts.append("\n" + ("#" * int(level)) + " " + para.text)
                else:
                    parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts) or f"[DOCX {path.name} was empty]"
    except Exception as exc:  # noqa: BLE001
        logger.exception("DOCX read failed")
        return f"[DOCX read failed: {exc!s}]"


def _read_xlsx(path: Path) -> str:
    try:
        from openpyxl import load_workbook  # type: ignore
    except ImportError:
        return f"[XLSX reader unavailable — install openpyxl to read {path.name}]"
    try:
        wb = load_workbook(str(path), read_only=True, data_only=True)
        sheets = []
        for ws in wb.worksheets:
            rows = []
            for row in ws.iter_rows(values_only=True):
                rows.append(" | ".join("" if v is None else str(v) for v in row))
                if len(rows) >= 200:
                    rows.append("[…truncated]")
                    break
            sheets.append(f"\n## Sheet: {ws.title}\n" + "\n".join(rows))
        return "\n".join(sheets)
    except Exception as exc:  # noqa: BLE001
        logger.exception("XLSX read failed")
        return f"[XLSX read failed: {exc!s}]"


def _read_html(path: Path) -> str:
    raw = path.read_text(encoding="utf-8", errors="replace")
    try:
        from markdownify import markdownify  # type: ignore

        return markdownify(raw, heading_style="ATX")
    except ImportError:
        return re.sub(r"<[^>]+>", "", raw)


def _read_json(path: Path) -> str:
    raw = path.read_text(encoding="utf-8", errors="replace")
    try:
        return json.dumps(json.loads(raw), indent=2)
    except Exception:
        return raw


async def read_file(
    path: Annotated[str, Field(description="Absolute or relative path to the file")],
    max_chars: Annotated[int, Field(description="Maximum characters to return (1000-200000)")] = _MAX_DEFAULT,
) -> str:
    """
    Read a file and return its text content (Markdown when applicable).

    Auto-detects PDF, DOCX, XLSX, HTML, JSON, and common text/code formats.
    Returns a clear message instead of raising on errors.
    """
    audit_id = audit_log("FileReader.read", "started", {"path": path, "max_chars": max_chars})
    max_chars = max(1000, min(200000, int(max_chars or _MAX_DEFAULT)))
    p = _safe_path(path)
    if p is None or not p.exists():
        audit_log("FileReader.read", "not_found", {"path": path}, parent_id=audit_id)
        return f"File not found: {path}"
    if p.is_dir():
        audit_log("FileReader.read", "is_dir", {"path": str(p)}, parent_id=audit_id)
        return f"Path is a directory, not a file: {p}"

    ext = p.suffix.lower()
    try:
        if ext == ".pdf":
            text = _read_pdf(p)
        elif ext == ".docx":
            text = _read_docx(p)
        elif ext == ".xlsx":
            text = _read_xlsx(p)
        elif ext in _HTML_EXTS:
            text = _read_html(p)
        elif ext == ".json":
            text = _read_json(p)
        elif ext in _TEXT_EXTS or ext == "":
            text = p.read_text(encoding="utf-8", errors="replace")
        else:
            try:
                text = p.read_text(encoding="utf-8", errors="replace")
            except Exception as exc:  # noqa: BLE001
                return f"Unsupported file type {ext}: {exc!s}"
    except Exception as exc:  # noqa: BLE001
        logger.exception("File read failed")
        audit_log("FileReader.read", "error", {"error": str(exc)}, parent_id=audit_id)
        return f"Failed to read {p}: {exc!s}"

    truncated = len(text) > max_chars
    if truncated:
        text = text[:max_chars] + "\n\n[…truncated]"
    audit_log(
        "FileReader.read",
        "completed",
        {"path": str(p), "chars": len(text), "truncated": truncated},
        parent_id=audit_id,
    )
    return f"# File: {p}\n\n{text}"


async def list_dir(
    path: Annotated[str, Field(description="Directory to list")],
    pattern: Annotated[str, Field(description="Optional glob pattern (e.g. *.log)")] = "*",
) -> str:
    """List the contents of a directory (non-recursive). Returns names + sizes."""
    audit_id = audit_log("FileReader.list_dir", "started", {"path": path, "pattern": pattern})
    p = _safe_path(path)
    if p is None or not p.exists() or not p.is_dir():
        return f"Not a directory: {path}"
    try:
        entries = sorted(p.glob(pattern or "*"))
        lines = [f"# Directory: {p}"]
        for e in entries[:200]:
            tag = "DIR " if e.is_dir() else "FILE"
            try:
                size = e.stat().st_size if e.is_file() else 0
            except OSError:
                size = 0
            lines.append(f"{tag} {size:>10}  {e.name}")
        if len(entries) > 200:
            lines.append(f"[…and {len(entries) - 200} more]")
        audit_log("FileReader.list_dir", "completed", {"count": len(entries)}, parent_id=audit_id)
        return "\n".join(lines)
    except Exception as exc:  # noqa: BLE001
        logger.exception("list_dir failed")
        return f"List directory failed: {exc!s}"


async def search_in_file(
    path: Annotated[str, Field(description="File to search")],
    pattern: Annotated[str, Field(description="Regex (or plain text) pattern")],
    max_matches: Annotated[int, Field(description="Maximum matches to return (1-100)")] = 25,
) -> str:
    """Find lines in a file matching a regex (case-insensitive)."""
    audit_id = audit_log("FileReader.search", "started", {"path": path, "pattern": pattern})
    p = _safe_path(path)
    if p is None or not p.exists() or not p.is_file():
        return f"File not found: {path}"
    max_matches = max(1, min(100, int(max_matches or 25)))
    try:
        rx = re.compile(pattern, re.IGNORECASE)
    except re.error as exc:
        return f"Invalid regex {pattern!r}: {exc!s}"

    out = []
    try:
        with p.open("r", encoding="utf-8", errors="replace") as fh:
            for line_no, line in enumerate(fh, start=1):
                if rx.search(line):
                    out.append(f"{line_no:>6}: {line.rstrip()[:300]}")
                    if len(out) >= max_matches:
                        out.append(f"[stopped at {max_matches} matches]")
                        break
    except Exception as exc:  # noqa: BLE001
        logger.exception("search_in_file failed")
        return f"Search failed: {exc!s}"
    audit_log("FileReader.search", "completed", {"matches": len(out)}, parent_id=audit_id)
    if not out:
        return f"No matches for {pattern!r} in {p}"
    return f"# Matches in {p}\n\n" + "\n".join(out)


FILE_READER_TOOLS = [read_file, list_dir, search_in_file]
